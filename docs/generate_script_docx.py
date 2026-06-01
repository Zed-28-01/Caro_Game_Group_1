"""Generate presentation script as .docx for Google Docs import."""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font for the document
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0x70)
    return h

def add_para(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_script(text):
    """Add a script paragraph in italic indented (for spoken parts)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# === TITLE ===
title = doc.add_heading("SCRIPT THUYẾT TRÌNH — CARO GAME (20 phút)", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x70)

add_para("Nhóm 1 — Đồ án Cơ sở lập trình C++ | Presenter: Nhật", italic=True, color=(0x66, 0x66, 0x66))
doc.add_paragraph()

# === TIME BUDGET ===
add_heading("⏱️ Time Budget 20 phút", level=1)
budget_table = doc.add_table(rows=8, cols=3)
budget_table.style = "Light Grid Accent 1"
headers = budget_table.rows[0].cells
headers[0].text = "Phần"
headers[1].text = "Slide"
headers[2].text = "Thời gian"
budget_data = [
    ("1. Mở đầu", "1-3", "1:00"),
    ("2. Kiến trúc + cơ bản", "4-9", "4:30"),
    ("3. UI showcase", "10-15", "1:30"),
    ("4. Kỹ thuật nâng cao", "16-21", "5:00"),
    ("5. DEMO ⭐", "22", "6:00"),
    ("6. Tổng kết", "23-25", "1:30"),
    ("Buffer Q&A", "—", "0:30"),
]
for i, (a, b, c) in enumerate(budget_data, 1):
    row = budget_table.rows[i].cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()

# ============================================================
# PHẦN 1 — MỞ ĐẦU
# ============================================================
add_heading("🎬 PHẦN 1 — MỞ ĐẦU (1 phút)", level=1)

add_heading("Slide 1 — Title (10s)", level=2)
add_script(
    "Chào thầy và các bạn. Em là Nhật, đại diện Nhóm 1 trình bày đồ án "
    "Caro Game — game cờ caro 5 quân thắng, phát triển bằng C++ và thư viện SFML. "
    "Buổi thuyết trình hôm nay sẽ kéo dài khoảng 20 phút, bao gồm phần demo trực tiếp 6 phút."
)

add_heading("Slide 2 — Nội dung trình bày (15s)", level=2)
add_script(
    "Nội dung em sẽ trình bày gồm 4 phần: thứ nhất, kiến trúc tổng thể và các tính năng "
    "cơ bản theo đề bài; thứ hai, các kỹ thuật nâng cao như Bot AI và Speed Mode; thứ ba, "
    "demo trực tiếp sản phẩm; và cuối cùng là tổng kết những gì nhóm em đã học được."
)

add_heading("Slide 3 — Giới thiệu game (35s)", level=2)
add_script(
    "Caro Game hay còn gọi là Gomoku — luật chơi đơn giản: ai đặt được 5 quân liên tiếp "
    "trên bàn cờ 15×15 thì thắng. Nhóm em đã hiện thực hóa game này thành sản phẩm khoảng "
    "3000 dòng C++, chia thành 6 module rõ ràng. Phần graphics và audio em dùng thư viện "
    "SFML 2.6.2. Toàn bộ quá trình phát triển nhóm em quản lý bằng Git và GitHub — đảm bảo "
    "code có lịch sử thay đổi rõ ràng, dễ phối hợp giữa các thành viên."
)
add_para("💡 Pacing tip: Đọc chậm phần \"5 quân liên tiếp\", \"15×15\", \"3000 dòng\", \"6 module\".",
         italic=True, color=(0x80, 0x60, 0x30), size=10)

# ============================================================
# PHẦN 2 — KIẾN TRÚC + CƠ BẢN
# ============================================================
add_heading("🏗️ PHẦN 2 — KIẾN TRÚC + CƠ BẢN (4 phút 30 giây)", level=1)

add_heading("Slide 4 — Cấu trúc module (50s)", level=2)
add_script(
    "Đây là sơ đồ phụ thuộc của 9 module — chỉ ra cách các file source liên kết với nhau."
)
add_script(
    "Trung tâm là game_types.h — chứa toàn bộ struct, enum, constant. Mọi module khác đều "
    "include file này, đóng vai trò \"từ vựng chung\"."
)
add_script(
    "main.cpp là điểm vào, tạo window và gọi gameRun() trong menu.cpp. menu.cpp là dispatcher "
    "chính — điều phối luồng giữa các màn hình, gọi 6 module chức năng: bot, board, timer, "
    "render, save_load, sound, language."
)
add_script(
    "Kiến trúc phẳng kiểu C, không OOP nặng — phù hợp với class project, dễ hiểu, dễ debug."
)

add_heading("Slide 5 — Sơ đồ luồng game (50s)", level=2)
add_script(
    "Đây là state diagram thể hiện luồng chuyển đổi giữa các màn hình. Bắt đầu từ Main Menu, "
    "người chơi có thể vào game mới qua 4 bước: chọn Mode (PvP/PvC), chọn độ khó (nếu PvC), "
    "chọn Style (Basic/Speed), nhập tên — rồi vào Playing."
)
add_script(
    "Trong Playing có 3 đường thoát: nhấn ESC vào Pause Menu, nhấn L để Save, hoặc thắng/hòa/"
    "hết giờ vào Game Over."
)
add_script(
    "Tính năng mới nhóm em vừa thêm: ở Game Over, sau khi hỏi \"Chơi tiếp?\" và người chơi "
    "chọn Không, em hỏi tiếp \"Lưu ván đấu?\" — giúp người chơi không mất ván hay."
)

add_heading("Slide 6 — Cấu trúc dữ liệu (40s)", level=2)
add_script("GameState là god struct chứa toàn bộ trạng thái 1 ván cờ. Có vài chi tiết đáng lưu ý:")
add_script("Thứ nhất, board là mảng tĩnh 2D 15×15 — nhanh hơn vector, cache friendly.")
add_script(
    "Thứ hai, moveHistory là mảng tĩnh 225 phần tử + moveCount — em dùng pattern array+count "
    "thay cho std::stack để serialize save/load trivially."
)
add_script(
    "Thứ ba, TimerState em dùng chess-clock — mỗi người có 10 phút riêng, sẽ trình bày sau."
)

add_heading("Slide 7 — Kiểm tra thắng (50s)", level=2)
add_script(
    "Thay vì quét toàn bộ 225 ô mỗi lượt, em chỉ check từ ô vừa đặt theo 4 hướng: ngang, dọc, "
    "2 chéo."
)
add_script(
    "Mỗi hướng em đếm về cả 2 phía — phía dương và phía âm — rồi cộng với ô gốc. Nếu tổng ≥ 5 "
    "thì thắng. Em chỉ check sau mỗi nước đi, độ phức tạp O(1) mỗi lượt thay vì O(N²) — nhanh "
    "gấp ~225 lần."
)

add_heading("Slide 8 — Kiểm tra hòa (20s)", level=2)
add_script(
    "Tương tự, em tối ưu hàm checkDraw xuống O(1) bằng cách dùng biến moveCount — luôn đồng "
    "bộ với số ô đã đặt qua các thao tác đặt/undo/reset. Bàn đầy khi moveCount = 225. Chỉ 1 "
    "phép so sánh int."
)

add_heading("Slide 9 — Save/Load (50s)", level=2)
add_script(
    "Save/Load em dùng định dạng text thay vì binary. Lý do: text dễ debug, dễ hỗ trợ UTF-8 "
    "cho tên người chơi tiếng Việt, và không phụ thuộc layout struct."
)
add_script(
    "Quá trình save em ghi tuần tự thông tin 2 player, trạng thái ván, bàn cờ, lịch sử nước "
    "đi. Load thì đọc ngược lại theo đúng thứ tự đó."
)
add_script("Tổng cộng game hỗ trợ tối đa 15 slot save.")

# ============================================================
# PHẦN 3 — UI SHOWCASE
# ============================================================
add_heading("🎨 PHẦN 3 — UI SHOWCASE (1 phút 30 giây)", level=1)
add_para("Lướt nhanh 6 screenshots, mỗi slide ~15 giây.", italic=True, color=(0x66, 0x66, 0x66))

add_heading("Slide 10 — Main Menu (15s)", level=2)
add_script(
    "Màn hình chính: logo CARO sticker, 6 menu items với button frame bo tròn vàng kim, "
    "background ảnh anime Dragon Ball xuyên suốt mọi màn hình."
)

add_heading("Slide 11 — Mode Select (15s)", level=2)
add_script("Chọn PvP hai người hoặc PvC chơi với máy.")

add_heading("Slide 12 — Difficulty (15s)", level=2)
add_script("4 cấp độ Bot: Easy, Medium, Hard, Expert — em sẽ giải thích chi tiết ở phần sau.")

add_heading("Slide 13 — Gameplay (20s)", level=2)
add_script(
    "Đây là màn hình chính: bàn cờ 15×15 ở giữa, panel 2 người chơi bên phải với mascot Goku "
    "và Vegeta. Trong Speed mode có thanh đếm ngược 20s/lượt và thời gian ván của mỗi người "
    "riêng biệt."
)

add_heading("Slide 14 — Pause Menu (15s)", level=2)
add_script("Nhấn ESC mở Pause Menu với 3 lựa chọn: tiếp tục, lưu, về Menu chính.")

add_heading("Slide 15 — Game Over (15s)", level=2)
add_script(
    "Khi kết thúc, banner VICTORY hiện cho người thắng, DEFEAT cho người thua. Mascot tự đổi "
    "tư thế tương ứng. Hỏi \"Chơi tiếp?\" rồi nếu Không thì hỏi \"Lưu ván?\"."
)

# ============================================================
# PHẦN 4 — KỸ THUẬT NÂNG CAO
# ============================================================
add_heading("🤖 PHẦN 4 — KỸ THUẬT NÂNG CAO (5 phút)", level=1)

add_heading("Slide 16 — Bot AI 4 cấp độ (45s)", level=2)
add_script("Bot AI có 4 cấp độ tăng dần:")
add_bullet("Easy: random với check thắng/chặn ngay — dành cho người mới làm quen")
add_bullet("Medium: heuristic scoring — nhìn xa 1 lượt")
add_bullet("Hard: Minimax + Alpha-Beta pruning, độ sâu 3")
add_bullet("Expert: tương tự Hard nhưng độ sâu 4 với Move Ordering")
add_script(
    "Cả Hard và Expert đều dùng Minimax kết hợp Alpha-Beta. Khác biệt là Expert nhìn sâu hơn "
    "và có cơ chế sắp xếp nước đi để cắt tỉa hiệu quả hơn."
)
add_script(
    "Tính năng Hint — phím H — dùng chính thuật toán Medium để gợi ý cho người chơi. Đủ thông "
    "minh nhưng đủ nhanh để phản hồi tức thì."
)

add_heading("Slide 17 — botGetCandidates (45s)", level=2)
add_script(
    "Trước khi chấm điểm, em lọc danh sách ô cần xét. Thay vì xét cả 225 ô, em chỉ xét ô "
    "trống có quân nào đó trong bán kính radius."
)
add_script(
    "Lý do: nước đi tiềm năng chỉ có nghĩa khi gần quân hiện có — đặt ngoài rìa khi chưa có "
    "gì là vô nghĩa."
)
add_script(
    "Trường hợp đặc biệt: bàn cờ trống → em trả về ô trung tâm (7,7) làm nước khai cuộc."
)
add_script(
    "Kết quả: candidates giảm từ 225 xuống 20-50 → bot Expert chạy được ở depth 4 trong thời "
    "gian chấp nhận được."
)

add_heading("Slide 18 — Pattern Scoring ⭐ (75s)", level=2)
add_script("Đây là building block cho 3 cấp bot có heuristic: Medium, Hard, Expert.")
add_script(
    "Hàm scoreLine xét 1 chuỗi quân theo 1 hướng. Đếm số quân liên tiếp về 2 phía, cộng với "
    "ô đang xét. Đồng thời đếm số đầu mở — tức là 2 đầu chuỗi có phải ô trống không."
)
add_script(
    "Rồi tra bảng điểm: 5 liên tiếp đương nhiên thắng — 1 triệu điểm. 4 mở 2 đầu chắc thắng "
    "vì không thể chặn cả 2 đầu trong 1 lượt — 100 ngàn điểm. 4 mở 1 đầu có thể chặn được — "
    "5 ngàn điểm. Tương tự cho 3, 2, 1 quân."
)
add_script(
    "Bậc thang điểm này em thiết kế để bot phân biệt rõ giữa các nước — chênh nhau 10-20 lần "
    "để bot không bị nhầm giữa nước tốt vừa và nước cực tốt."
)
add_script(
    "Medium dùng trực tiếp scoreLine để chấm. Hard và Expert dùng nó tại lá của cây Minimax "
    "— slide tiếp theo."
)

add_heading("Slide 19 — Minimax + Alpha-Beta ⭐⭐ ĐIỂM NHẤN (2 phút)", level=2)
add_para("⚠️ ĐOẠN QUAN TRỌNG NHẤT — đọc chậm + nhấn mạnh tác giả gốc.", bold=True, color=(0xC0, 0x39, 0x2B))
add_script("Đây là điểm nhấn của phần Bot AI.")
add_script(
    "Minimax là thuật toán cổ điển trong game theory — được John von Neumann chứng minh năm "
    "1928, Claude Shannon áp dụng vào AI cờ vua năm 1950, và IBM Deep Blue dùng để đánh bại "
    "Garry Kasparov năm 1997. Nhóm em không phát minh thuật toán này — em triển khai lại cho "
    "game Caro."
)
add_script(
    "Ý tưởng: Bot giả định cả 2 người chơi đều tối ưu. Bot là MAX — tìm nước có điểm cao "
    "nhất. Đối thủ là MIN — bot giả định đối thủ sẽ chọn nước làm bot bất lợi nhất. Đệ quy "
    "này lan truyền đến độ sâu maxDepth, đến lá em gọi evaluateBoard đánh giá toàn cảnh bằng "
    "scoreLine."
)
add_script(
    "Vấn đề: depth 4 phải duyệt rất nhiều trạng thái. Alpha-Beta giải quyết: nếu một nhánh đã "
    "chắc chắn tệ hơn nhánh tốt nhất đã tìm — cắt luôn không duyệt nữa. Cụ thể: beta ≤ alpha "
    "thì break."
)
add_script(
    "Đóng góp riêng của nhóm em là: thiết kế bảng điểm pattern phù hợp riêng cho Caro, chọn "
    "depth 4 cân bằng giữa thông minh và tốc độ — bot phản hồi dưới 200ms — và tích hợp với "
    "chess-clock của Speed mode."
)

add_heading("Slide 20 — Speed Mode Chess-Clock (50s)", level=2)
add_script(
    "Speed Mode em thiết kế theo cơ chế chess-clock — mỗi người có 10 phút riêng, đồng hồ chỉ "
    "chạy khi đến lượt mình."
)
add_script(
    "Mỗi frame, game loop tính deltaTime từ sf::Clock, trừ vào thời gian lượt và thời gian "
    "ván của người đang đi. Có 3 điều kiện kết thúc: hết 20 giây của lượt — người đang đi "
    "thua; P1 hết 10 phút — P1 thua; P2 hết 10 phút — P2 thua."
)
add_script(
    "Quan trọng: nhóm em xử lý cẩn thận trường hợp PvC + Speed. Khi bot Expert suy nghĩ 2-5 "
    "giây, em đo bằng sf::Clock riêng và trừ thủ công vào thời gian của bot — đảm bảo "
    "chess-clock công bằng."
)

add_heading("Slide 21 — Mouse Support (30s)", level=2)
add_script(
    "Mouse Support em làm bằng cách chuyển đổi tọa độ pixel chuột về tọa độ ô trên bàn cờ."
)
add_script(
    "Hàm pixelToBoard chia tọa độ cho CELL_SIZE (40px) sau khi trừ offset bàn cờ — ra (row, "
    "col) từ 0-14."
)
add_script(
    "Trong game loop: di chuột làm cursor follow chuột, click trái thì đặt quân. Ngoài ra hàm "
    "menuHitTest xác định chuột đang nằm trên nút bấm nào trong menu để highlight."
)

# ============================================================
# PHẦN 5 — DEMO
# ============================================================
add_heading("🎮 PHẦN 5 — DEMO TRỰC TIẾP (6 phút) ⭐ CRITICAL", level=1)

add_heading("Mở đầu demo", level=2)
add_script("Bây giờ em xin phép chuyển sang demo trực tiếp sản phẩm.")

add_heading("Phần 1 — Main Menu + Settings (45s)", level=2)
add_para("Các bước thực hiện:", bold=True)
add_bullet("Mở exe → Main Menu hiện ra")
add_bullet("Vào Settings → kéo slider volume → đổi VN/EN")
add_bullet("ESC về Main")
add_para("Câu nói:", bold=True)
add_script("Đây là màn hình chính. Em show qua menu Cài đặt...")
add_script(
    "Game hỗ trợ tiếng Việt và tiếng Anh, settings được lưu vào file saves/settings.txt."
)

add_heading("Phần 2 — PvP Basic (1 phút)", level=2)
add_para("Các bước thực hiện:", bold=True)
add_bullet("Chọn: Chơi mới → PvP → Basic → nhập tên 2 player tiếng Việt")
add_bullet("Đặt vài quân với WASD + Enter, cố ý tạo chuỗi 4 quân")
add_bullet("Tiếp tục đến khi thắng → highlight 5 quân thắng màu vàng")
add_bullet("Chọn \"Không\" ở \"Chơi tiếp?\" → hiện câu hỏi mới về lưu ván")
add_bullet("Chọn \"Không\" → về Main Menu")
add_para("Câu nói:", bold=True)
add_script("Em sẽ chơi nhanh 1 ván PvP để show gameplay và mascot.")
add_script(
    "Cursor di bằng WASD, đặt quân bằng Enter, hoặc dùng chuột — em sẽ show chuột ở ván sau."
)
add_script(
    "5 quân thắng được highlight, mascot Goku đổi tư thế chiến thắng, banner VICTORY hiện ra."
)
add_script("Đây là tính năng nhóm em vừa thêm: hỏi có muốn lưu ván đấu không.")

add_heading("Phần 3 — PvC Expert + Speed (2 phút 30 giây) ⭐ ĐIỂM NHẤN DEMO", level=2)
add_para("Các bước thực hiện:", bold=True)
add_bullet("Chọn: Chơi mới → PvC → Expert → Speed → nhập tên")
add_bullet("Đặt quân vài lần — show cả bàn phím và chuột")
add_bullet("Cố ý nhấn H để gọi Hint → highlight ô gợi ý")
add_bullet("Tiếp tục đánh, để bot thắng nếu có thể (Expert thường thắng)")
add_para("Câu nói:", bold=True)
add_script("Bây giờ em đánh với bot cấp Expert ở chế độ Speed.")
add_script(
    "Mỗi người có 10 phút riêng, hiển thị trong panel. Khi đến lượt mình, đồng hồ chạy. "
    "Khi bot suy nghĩ, có text \"Bot đang suy nghĩ\" báo cho người dùng biết."
)
add_script("Phím H gọi gợi ý — dùng thuật toán Medium đánh giá theo góc nhìn người chơi.")
add_script("Bot Expert nhìn xa 4 lượt và sort candidates trước khi search — phản hồi vẫn dưới 200ms.")

add_heading("Phần 4 — Save/Load (1 phút)", level=2)
add_para("Các bước thực hiện:", bold=True)
add_bullet("Chơi tiếp → Chọn Có → reset round (loser starts)")
add_bullet("Đặt 5-6 quân, nhấn L → Save Screen")
add_bullet("Nhập tên save \"demo_test\" → lưu")
add_bullet("ESC về Playing → ESC nữa → Pause Menu → Về Menu")
add_bullet("Main Menu → Tải game → chọn \"demo_test\"")
add_para("Câu nói:", bold=True)
add_script("Theo quy tắc cờ vua truyền thống, người thua ván trước đi trước ván sau.")
add_script("Save lưu dưới định dạng text trong thư mục saves.")
add_script(
    "Load thành công — toàn bộ trạng thái khôi phục đúng: vị trí quân, thời gian từng người, "
    "lượt đang đi."
)

add_heading("Phần 5 — Resize window (30s)", level=2)
add_para("Các bước thực hiện:", bold=True)
add_bullet("Kéo to window hoặc nhấn nút Maximize")
add_bullet("Chuột vẫn click đúng ô")
add_para("Câu nói:", bold=True)
add_script(
    "Game hỗ trợ resize giữ tỷ lệ 16:9 bằng letterbox — chi tiết kỹ thuật em có ghi trong báo cáo."
)
add_script("Mouse hit-test vẫn chính xác sau khi resize.")

add_heading("Kết thúc demo (15s)", level=2)
add_bullet("ESC về Main Menu")
add_script("Đó là toàn bộ demo. Em xin trở lại slide tổng kết.")

# ============================================================
# PHẦN 6 — TỔNG KẾT
# ============================================================
add_heading("🏁 PHẦN 6 — TỔNG KẾT (1 phút 30 giây)", level=1)

add_heading("Slide 23 — Đối chiếu yêu cầu cốt lõi (45s)", level=2)
add_script("Em đối chiếu với 5 yêu cầu cốt lõi của đề bài:")
add_bullet("Save/Load: vượt yêu cầu với 15 slot")
add_bullet("Nhận biết thắng/thua/hòa: tối ưu O(1) mỗi lượt")
add_bullet("Hiệu ứng: có animation cho đặt quân, win, draw + sound effects")
add_bullet("Giao diện gameplay: có panel stats đầy đủ — tên, số bước, thắng, thời gian")
add_bullet("Màn hình chính: 13 màn hình tổng cộng")
add_script("Tất cả 5 yêu cầu cốt lõi đạt mức vượt.")

add_heading("Slide 24 — Đối chiếu yêu cầu mở rộng (30s)", level=2)
add_script("Ngoài ra nhóm em sáng tạo thêm:")
add_bullet("Bot AI 4 cấp + Hint + Undo")
add_bullet("Speed Mode với chess-clock")
add_bullet("Mouse Support")
add_bullet("Đa ngôn ngữ Việt-Anh và lưu settings")
add_bullet("Một số tính năng âm thanh, UI/UX mượt cho game")

add_heading("Slide 25 — Tổng kết (15s)", level=2)
add_script(
    "Tổng kết — kỹ thuật em đã áp dụng: mảng 2D, struct, đệ quy, thiết kế kiểu pointer-free. "
    "Học được: SFML, Minimax, Alpha-Beta, Game Loop, File I/O. Khó khăn lớn nhất là đồng bộ "
    "UI 13 màn hình và tối ưu bot dưới 200ms phản hồi."
)
add_script(
    "Em xin kết thúc phần trình bày. Cảm ơn thầy và các bạn đã lắng nghe. Em sẵn sàng nhận "
    "câu hỏi từ Hội đồng."
)

# ============================================================
# PACING CHEATSHEET
# ============================================================
doc.add_page_break()
add_heading("🎯 PACING CHEATSHEET (in sẵn liếc khi present)", level=1)

pacing_table = doc.add_table(rows=7, cols=3)
pacing_table.style = "Light Grid Accent 1"
pacing_table.rows[0].cells[0].text = "Mốc thời gian"
pacing_table.rows[0].cells[1].text = "Đang ở slide nào"
pacing_table.rows[0].cells[2].text = "Nếu lệch"
pacing_data = [
    ("02:00", "Hết slide 4-5", "Lệch >30s → rút gọn UI showcase"),
    ("05:30", "Hết slide 9", "Lệch >1min → bỏ slide 14 (Pause)"),
    ("07:00", "Hết slide 15", "Đang đúng tiến độ"),
    ("12:00", "Hết slide 21 (chuyển demo)", "Demo 6 phút từ đây"),
    ("18:00", "Hết demo (slide 22)", "Còn 2 phút cho 3 slide kết"),
    ("19:30", "Hết slide 25", "OK — buffer Q&A 30s"),
]
for i, (a, b, c) in enumerate(pacing_data, 1):
    row = pacing_table.rows[i].cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()

# ============================================================
# TIPS DEMO
# ============================================================
add_heading("💡 TIPS DEMO CRITICAL", level=1)
add_bullet("Trước demo: Xóa saves/Gamelist.txt (do format thay đổi) — chỉ giữ settings.txt")
add_bullet("Chuẩn bị sẵn 1 save game test trước để show Load nhanh")
add_bullet("Tập demo 3 lần trước thuyết trình thật — đếm giờ")

add_para("Phím tắt cần nhớ:", bold=True)
add_bullet("W/A/S/D: di cursor")
add_bullet("Enter: đặt quân")
add_bullet("H: hint (chỉ PvC)")
add_bullet("L: save")
add_bullet("Z: undo")
add_bullet("ESC: pause")

add_para("Backup plan:", bold=True)
add_bullet("Nếu game crash: \"Em có quay video demo dự phòng\" (chuẩn bị sẵn nếu có thể)")
add_bullet("Nếu lệch thời gian: cắt UI showcase (slide 14 Pause + 1 slide khác)")

# ============================================================
# Q&A FLASHCARDS QUICK REFERENCE
# ============================================================
doc.add_page_break()
add_heading("📚 Q&A FLASHCARDS — TOP 10 CÂU HỎI THƯỜNG GẶP", level=1)

qa = [
    ("Minimax là của các em phát minh à?",
     "Không. John von Neumann chứng minh năm 1928, Claude Shannon áp dụng AI cờ vua 1950. Em TRIỂN KHAI lại cho Caro. Đóng góp riêng: bảng điểm pattern + depth=4 + tích hợp chess-clock."),
    ("Tại sao chọn Minimax mà không phải MCTS?",
     "Caro là zero-sum + deterministic + perfect-info → Minimax tối ưu lý thuyết. Branching factor nhỏ hơn cờ vua → depth 4 đủ thông minh. Có scoreLine chính xác → không cần simulation như MCTS."),
    ("Hard và Expert khác nhau gì?",
     "Cả 2 đều Minimax + Alpha-Beta. Khác: Hard depth 3 / Expert depth 4 + Move Ordering (sort candidates) + Top-12 filter."),
    ("evaluatePosition vs evaluateBoard?",
     "evaluatePosition chấm 1 nước đi (mô phỏng đặt quân, attack + 0.9×defense + center bonus). Dùng cho Medium, Move Ordering Expert, Hint. evaluateBoard chấm toàn bàn cờ. Dùng tại lá Minimax (Hard + Expert)."),
    ("Tại sao defense = 0.9 × attack?",
     "Em test 3 tỷ lệ: 1.0 → bot quá phòng thủ. 0.5 → bot bỏ qua đối thủ. 0.9 → cân bằng — ưu tiên tấn công nhưng vẫn block."),
    ("Tại sao chess-clock thay vì shared timer?",
     "Công bằng — người chậm tự bị phạt, không phụ thuộc đối thủ. Ai hết time → người đó thua (giống cờ vua truyền thống)."),
    ("UTF-8 Vietnamese hiển thị thế nào?",
     "Dùng sf::String::fromUtf8(s.begin(), s.end()). Lý do: setString(std::string) interpret bytes là Latin-1 → mojibake với UTF-8."),
    ("Sao mouse hit-test vẫn đúng sau resize?",
     "handleCommonEvent map pixel chuột → view coord (1280x720 chuẩn) trước khi đến code hit-test. Letterbox giữ tỷ lệ 16:9."),
    ("Save/Load dùng binary hay text?",
     "Text. Lý do: dễ debug, UTF-8 friendly cho tên VN, không phụ thuộc struct layout."),
    ("Undo dùng std::stack hay mảng?",
     "Mảng tĩnh + count. Lý do: kích thước max biết chắc (225), cache locality, save/load đơn giản, random access cho PvC undo 2 nước."),
]

for q, a in qa:
    p = doc.add_paragraph()
    run = p.add_run("Q: " + q)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    add_script("A: " + a)
    doc.add_paragraph()

# Save
out_path = r"D:\HCMUS\Programming C++\TTT\Caro_Project\Caro_Game_Group_1\docs\PRESENTATION_SCRIPT.docx"
doc.save(out_path)
print(f"DONE: {out_path}")
